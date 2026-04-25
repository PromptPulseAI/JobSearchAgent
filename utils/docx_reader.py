"""
Read .docx files and extract text or structured sections.
Uses python-docx. All extracted text is plain UTF-8 strings — no binary output.
"""
from pathlib import Path
from typing import Dict, Union

import docx

from utils.exceptions import FileIOError
from utils.logger import audit


def read_docx_text(filepath: Union[str, Path], agent: str = "system") -> str:
    """
    Extract full plain text from a .docx file.
    Paragraphs are joined with newlines; empty paragraphs are skipped.
    """
    path = Path(filepath)
    _check_path(path, agent)
    try:
        doc = docx.Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        text = "\n".join(paragraphs)
        audit("read", agent, "docx", "success", detail=f"{len(paragraphs)} paragraphs from {path.name}")
        return text
    except Exception as exc:
        audit("read", agent, "docx", "failure", detail=str(exc))
        raise FileIOError(f"Failed to read {path}: {exc}", agent=agent)


def read_docx_sections(filepath: Union[str, Path], agent: str = "system") -> Dict[str, str]:
    """
    Extract named sections from a .docx file.
    A paragraph is treated as a section header if it uses a Heading style,
    is bold throughout, or is all-uppercase and short (<= 50 chars).
    Returns {"section_name_lower": "section body text", ...}.
    """
    path = Path(filepath)
    _check_path(path, agent)
    try:
        doc = docx.Document(str(path))
    except Exception as exc:
        raise FileIOError(f"Failed to open {path}: {exc}", agent=agent)

    sections: Dict[str, str] = {}
    current_key = "preamble"
    current_lines = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue

        if _is_header(para):
            if current_lines:
                sections[current_key] = "\n".join(current_lines)
            current_key = text.lower()
            current_lines = []
        else:
            current_lines.append(text)

    if current_lines:
        sections[current_key] = "\n".join(current_lines)

    audit("read", agent, "docx_sections", "success", detail=f"{len(sections)} sections from {path.name}")
    return sections


def validate_docx(filepath: Union[str, Path]) -> bool:
    """Return True if the file is a valid, non-empty .docx. Used after generation."""
    try:
        doc = docx.Document(str(filepath))
        return len(doc.paragraphs) > 0
    except Exception:
        return False


def _check_path(path: Path, agent: str) -> None:
    if not path.exists():
        raise FileIOError(f"File not found: {path}", agent=agent)
    if path.suffix.lower() != ".docx":
        raise FileIOError(f"Expected .docx file, got: {path.suffix}", agent=agent)


def _is_header(para) -> bool:
    """Heuristic to detect section headers in a .docx paragraph."""
    text = para.text.strip()
    if para.style.name.startswith("Heading"):
        return True
    if text.isupper() and 2 < len(text) <= 50:
        return True
    runs_with_text = [r for r in para.runs if r.text.strip()]
    if runs_with_text and all(r.bold for r in runs_with_text):
        return True
    return False
